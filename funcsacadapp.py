
import pypdf
from subprocess import call
import docx 
import os
from datetime import datetime, date
from lxml import etree
import zipfile
import docx2pdf
import aspose.words as aw



"""
ANNOTATION REPORT FUNCS
"""

def annotation_report_pdf(source: str):

    try:
        # reads the file
        input1 = pypdf.PdfReader(open(source, "rb"))
        nPages = input1.get_num_pages()
    except pypdf.errors.PdfReadError:
        return 'File cannot be read and is potentially corrupted', 409

    doc = docx.Document()                   # sets up the word document
    doc.add_heading("Annotation summary:", 1) # creates a title for doc
    doc.add_paragraph().add_run("")
    #doc.add_heading("Paper title: \n" + titleText.splitlines()[0] + titleText.splitlines()[1], 2)

    # loops through each page
    for i in range(nPages) :
        # get the data/text from this PDF page plus annotations
        page = input1.get_page(i)
        text = page.extract_text()

        # this contains the page title and comments of each page
        reportPageContent = []

        try :

            if page['/Annots']:

                # the page number of the page
                pageNumber = 'Page: ' + str(i + 1)  
                # adds page num as first element
                reportPageContent.append(pageNumber)

                for annot in page['/Annots'] :

                    # Other subtypes, such as /Link, cause errors
                    subtype = annot['/Subtype']
                    if subtype == "/Text":

                        newAnnotation = "- " + annot['/Contents']

                        reportPageContent.append(newAnnotation)
                
                # if no comments on page or only empty ones then pass, else add page and comments to report
                if len(reportPageContent) > 1:

                    pageTitle = reportPageContent[0]
                    doc.add_heading(pageTitle, 4)

                    for comment in range(1, len(reportPageContent)):

                        annotFromList = reportPageContent[comment]

                        doc.add_paragraph().add_run(annotFromList)

                else:
                    pass
                        
        except KeyError:
            # if no annotations on this page then move to next
            pass

    date = datetime.today().strftime('%Y-%m-%d')    # gets current date

    originalName = os.path.basename(source)[:-4]    # just removes the .pdf ending so it doesnt end up in the new word doc name

    worddocNewname = originalName + "_annotations_" + str(date) + ".docx"     # combines the source file name with "_annotations" (makes a name for the report)

    directory = os.path.split(source)[0]    # gets the directory path

    worddocPath = os.path.join(directory, worddocNewname)
        
    doc.save(worddocPath)

    return worddocPath


def annotation_report_word(wordPath: str, pdfPath: str, annots:str):

    try:
        readPDF = pypdf.PdfReader(open(pdfPath, "rb"))
        nPages = readPDF.get_num_pages()

        doc = docx.Document()                   # sets up the word document
        doc.add_heading("Annotation summary:", 1) # creates a title for doc
        doc.add_paragraph().add_run("")
    except pypdf.errors.PdfReadError:
        return 'File cannot be read and is potentially corrupted', 409
    
    # go through each page
    for pageIndex in range(nPages):

        firstcommentFound = False                   # is set to true after first comment on a page is found 

        page = readPDF.get_page(pageIndex)
        text = page.extract_text()
        
        try:
            for comment in annots:
                if text.find(comment) != -1 and type(comment) is not int:        # checks if comment is on page, if true, then it prints page num and comment if first on page, otherwise only comment
                    if firstcommentFound == False and comment.strip(" ") != "":
                        pageNumber = "Page: " + str(pageIndex + 1)
                        doc.add_heading(pageNumber, 4)
                        newAnnotation = "- " + comment
                        doc.add_paragraph().add_run(newAnnotation)
                        firstcommentFound = True
                    else:
                        newAnnotation = "- " + comment
                        doc.add_paragraph().add_run(newAnnotation)

        except TypeError:
            return 'no comments in the document', 400


    date = datetime.today().strftime('%Y-%m-%d')    # gets current date

    originalName = os.path.basename(wordPath)[:-5]    # just removes the .docx ending so it doesnt end up in the new word doc name

    worddocNewname = originalName + "_annotations_" + str(date) + ".docx"     # combines the source file name with "_annotations" (makes a name for the report)

    directory = os.path.split(wordPath)[0]    # gets the directory path

    worddocPath = os.path.join(directory, worddocNewname)
        
    doc.save(worddocPath)

    return worddocPath



def get_word_comments(source: str):

    if source.endswith(".docx") or source.endswith(".doc"):

        #os.chmod(source, stat.S_IRWXO )

        try:

            ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            annots = []

            docxZip = zipfile.ZipFile(source)
            commentsXML = docxZip.read('word/comments.xml')
            et = etree.XML(commentsXML)
            comments = et.xpath('//w:comment', namespaces=ooXMLns)

            for c in comments:
                # string value of the comment:
                annots.append(c.xpath('string(.)', namespaces=ooXMLns))

            if len(annots) != 0:
                return annots
            else:
                return 'There are no annotations in the selected file', 400
        except KeyError:
            return 'There are no annotations in the selected file', 400
    
    else:
        return 'Please input a valid file (Word (.docx) or PDF (.pdf))', 400
    

# nooo fix the path join here
def convert_docx_to_pdf(source: str):
   
    dirName = os.path.dirname(source)
    baseName = os.path.basename(source)
    
    pdfPath = dirName + "/" + baseName[:-5] + ".pdf"

    docx2pdf.convert(source, pdfPath)
    return pdfPath


def annotation_report_master_function(path_to_secure_file: str, basedir: str):

    try:

        if path_to_secure_file.endswith(".pdf"):
            wordDocpath = annotation_report_pdf(path_to_secure_file)
            # maybe add some user feedback here

        elif path_to_secure_file.endswith(".docx"):
            pdfPath = convert_docx_to_pdf(path_to_secure_file)
            comments = get_word_comments(path_to_secure_file)
            wordDocpath = annotation_report_word(path_to_secure_file, pdfPath, comments)
            # maybe add some user feedback here
        return wordDocpath
    
    # this is for when an invalid filetype is entered
    except UnboundLocalError:
        return 'Please select a valid filetype', 400
    


"""
DOC COMPARE FUNCS
"""

# fix comment deletion here
def compare_docx(uploadFolder: str, firstPath: str, secondPath: str):

    current_date = str(date.today())
    directory = uploadFolder

    # load first document
    doc = aw.Document(firstPath)

    # load second document
    doc2 = aw.Document(secondPath)

    ## FIX THIS SO IT MAKES A COPY WITHOUT COMMENTS AND KEEPS THE ORIGINAL WITH COMMENTS (USE WITHOUT FOR COMPARISON THEN DELETE, KEEPING THE ONE WITH COMMENTS)
    doc.revisions.accept_all()
    doc2.revisions.accept_all()

    # compare documents
    doc.compare(doc2, "Changes:", date.today())

    worddocName = os.path.basename(firstPath)[:-5] + "_comparison_" + current_date + ".docx"
    comparisonReportPath = os.path.join(directory, worddocName)
    # save the document to get the revisions

    #doc.range.replace("_Created with an evaluation copy of Aspose.Words. To remove all limitations, you can use Free Temporary License https://products.aspose.com/words/temporary-license/", 
    #                  "", aw.replacing.FindReplaceOptions(aw.replacing.FindReplaceDirection.FORWARD))

    
    doc.save(comparisonReportPath)

    return comparisonReportPath


def convert_pdf_to_docx(firstPDF: str, secondPDF: str):

    doc1 = docx.Document()
    doc2 = docx.Document()

    try:
        readpdf1 = pypdf.PdfReader(open(firstPDF, "rb"))
        nPages1 = readpdf1.get_num_pages()

        readpdf2 = pypdf.PdfReader(open(secondPDF, "rb"))
        nPages2 = readpdf2.get_num_pages()
    except pypdf.errors.PdfReadError:
        return 'File cannot be read and is potentially corrupted', 409


    for pageIndex in range(nPages1):

        page = readpdf1.get_page(pageIndex)
        text = page.extract_text()

        doc1.add_paragraph().add_run(text)


    for pageIndex in range(nPages2):

        page = readpdf2.get_page(pageIndex)
        text = page.extract_text()

        doc2.add_paragraph().add_run(text)


    originalName1 = os.path.basename(firstPDF)[:-4]    # just removes the .pdf ending so it doesnt end up in the new word doc name
    originalName2 = os.path.basename(secondPDF)[:-4]

    worddocNewname1 = originalName1 + ".docx" 
    worddocNewname2 = originalName2 + ".docx"                                      # combines the source file name with "_annotations" (makes a name for the report)

    directory1 = os.path.split(firstPDF)[0]    # gets the directory path
    directory2 = os.path.split(secondPDF)[0]

    worddocPath1 = os.path.join(directory1, worddocNewname1)
    worddocPath2 = os.path.join(directory2, worddocNewname2)
        
    doc1.save(worddocPath1)
    doc2.save(worddocPath2)

    return worddocPath1, worddocPath2


def convert_single_pdf_to_docx(PDFPath: str):

    if PDFPath.endswith(".pdf"):

        doc = docx.Document()

        try:
            readpdf = pypdf.PdfReader(open(PDFPath, "rb"))
            nPages = readpdf.get_num_pages()
        except pypdf.errors.PdfReadError:
            return 'File cannot be read and is potentially corrupted', 409


        for pageIndex in range(nPages):

            page = readpdf.get_page(pageIndex)
            text = page.extract_text()

            doc.add_paragraph().add_run(text)

        originalName1 = os.path.basename(PDFPath)[:-4]    # just removes the .pdf ending so it doesnt end up in the new word doc name

        worddocNewname = originalName1 + ".docx" 
                                        
        directory = os.path.split(PDFPath)[0]    # gets the directory path

        worddocPath = os.path.join(directory, worddocNewname)
            
        doc.save(worddocPath)

        return worddocPath


def document_comparison_master_function(uploadFolder: str, file1: str, file2:str):


    firstPath = os.path.join(uploadFolder, file1)  # gets the content of listbox with the path in it
    secondPath = os.path.join(uploadFolder, file2)



    if firstPath.endswith(".pdf") and secondPath.endswith(".pdf"):

                
        wordPaths = convert_pdf_to_docx(firstPath, secondPath)
        firstDocx = wordPaths[0]
        secondDocx = wordPaths[1]


        # save the two files here

        comparisonReportPath = compare_docx(uploadFolder, firstDocx, secondDocx)

        """
        if os.path.exists(firstDocx):
            os.remove(firstDocx)
        else:
            pass

        if os.path.exists(secondDocx):
            os.remove(secondDocx)
        else:
            pass
        """


    elif firstPath.endswith(".docx") and secondPath.endswith(".docx"):

                
        comparisonReportPath = compare_docx(uploadFolder, firstPath, secondPath)

       
    return comparisonReportPath

